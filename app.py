from flask import Flask, render_template, request, send_file
from docx import Document
from io import BytesIO
from xhtml2pdf import pisa
import uuid

app = Flask(__name__)

TEMPLATE_DOCX = "Nametag - Printable - Previous.docx"

def replace_in_paragraphs(paragraphs, replacements):
    for paragraph in paragraphs:
        for run in paragraph.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)

        # Fallback for split runs
        original_text = "".join(run.text for run in paragraph.runs)
        updated_text = original_text
        for old, new in replacements.items():
            updated_text = updated_text.replace(old, new)
        if paragraph.runs and updated_text != original_text:
            paragraph.runs[0].text = updated_text
            for run in paragraph.runs[1:]:
                run.text = ""

def replace_text(doc, replacements):
    replace_in_paragraphs(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, replacements)

def convert_html_to_pdf(source_html, output):
    pisa_status = pisa.CreatePDF(source_html, dest=output)
    return pisa_status.err

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # User inputs
        rank = request.form["rank"].upper()
        name_bn = request.form["name_bn"]
        name_en = request.form["name_en"].upper()
        pl = request.form["pl"].upper()
        serial = request.form["serial"].rstrip(".") + "."
        oc_no = request.form["oc_no"]

        replacements = {
            "OC": rank,
            "এহসান": name_bn,
            "EHSHAN": name_en,
            "M-4(A)": pl,
            "7.": serial,
            "14299": oc_no
        }

        # Load DOCX template
        doc = Document(TEMPLATE_DOCX)
        replace_text(doc, replacements)

        # Convert DOCX content to simple HTML
        html_content = ""
        for para in doc.paragraphs:
            html_content += f"<p>{para.text}</p>"

        # Add table content
        for table in doc.tables:
            html_content += "<table border='1' style='border-collapse: collapse;'>"
            for row in table.rows:
                html_content += "<tr>"
                for cell in row.cells:
                    html_content += f"<td style='padding:5px'>{cell.text}</td>"
                html_content += "</tr>"
            html_content += "</table>"

        # Generate PDF in memory
        pdf_file = BytesIO()
        error = convert_html_to_pdf(html_content, pdf_file)
        pdf_file.seek(0)

        if error:
            return "Error generating PDF", 500

        return send_file(
            pdf_file,
            as_attachment=True,
            download_name=f"{name_en} - Nametag.pdf"
        )

    return render_template("index.html")
